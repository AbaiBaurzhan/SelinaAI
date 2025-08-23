# rag.py
# RAG: PDF/DOCX/XLSX/JPG/PNG → чанки → эмбеддинги → поиск top-k + извлечение прайса

from __future__ import annotations

import os
import re
import json
import time
import base64
import sqlite3
from pathlib import Path
from typing import Callable, List, Tuple, Optional, Dict, Any

import numpy as np
import fitz  # PyMuPDF
from openai import OpenAI

from telegram import Update
from telegram.ext import ContextTypes
from telegram.constants import ChatAction

# DOCX/XLSX
from docx import Document as DocxDocument
from openpyxl import load_workbook

DB_PATH = Path(__file__).with_name("state.db")
UPLOADS_DIR = Path(__file__).parent / "uploads"
UPLOADS_DIR.mkdir(exist_ok=True)

# ----- резолвер ключа OpenAI (задаётся из main.py) -----
_KeyResolver: Optional[Callable[[], Optional[str]]] = None
def configure_key_resolver(func: Callable[[], Optional[str]]) -> None:
    global _KeyResolver
    _KeyResolver = func
def _resolve_api_key() -> Optional[str]:
    if _KeyResolver:
        key = _KeyResolver()
        if key:
            return key
    return os.getenv("OPENAI_API_KEY")

# ----- БД RAG -----
def db_init_rag() -> None:
    con = sqlite3.connect(DB_PATH)
    cur = con.cursor()
    cur.execute("""
        CREATE TABLE IF NOT EXISTS documents (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT,
            type TEXT,
            path TEXT,
            created_at INTEGER
        )
    """)
    cur.execute("""
        CREATE TABLE IF NOT EXISTS chunks (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            document_id INTEGER,
            idx INTEGER,
            text TEXT,
            embedding_json TEXT,
            FOREIGN KEY(document_id) REFERENCES documents(id)
        )
    """)
    # Новая таблица: нормализованные позиции прайса
    cur.execute("""
        CREATE TABLE IF NOT EXISTS catalog_items (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            document_id INTEGER,
            line_no INTEGER,
            name TEXT,
            price_value REAL,
            currency TEXT,
            raw_line TEXT,
            created_at INTEGER,
            FOREIGN KEY(document_id) REFERENCES documents(id)
        )
    """)
    con.commit()
    con.close()

# =======================
#        ПАРСЕРЫ
# =======================
def pdf_to_text(pdf_path: str) -> str:
    doc = fitz.open(pdf_path)
    texts: List[str] = []
    for page in doc:
        texts.append(page.get_text("text"))
    return "\n".join(texts)

def docx_to_text(path: str) -> str:
    doc = DocxDocument(path)
    parts: List[str] = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            line = " | ".join([c for c in cells if c])
            if line:
                parts.append(line)
    return "\n".join(parts)

def xlsx_to_text(path: str) -> str:
    wb = load_workbook(path, read_only=True, data_only=True)
    parts: List[str] = []
    for sheet in wb.worksheets:
        parts.append(f"[Лист: {sheet.title}]")
        for row in sheet.iter_rows(values_only=True):
            cells = [str(c).strip() for c in row if c is not None and str(c).strip() != ""]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)

# ===== OpenAI Vision: извлечение текста с изображений =====
def image_to_text_openai(image_path: str, api_key: str, model: str = "gpt-4o-mini") -> str:
    with open(image_path, "rb") as f:
        b64 = base64.b64encode(f.read()).decode("utf-8")
    lower = image_path.lower()
    if lower.endswith(".png"):
        mime = "image/png"
    elif lower.endswith(".jpg") or lower.endswith(".jpeg"):
        mime = "image/jpeg"
    else:
        mime = "image/jpeg"

    client = OpenAI(api_key=api_key)
    prompt = (
        "Извлеки весь текст с изображения прайса/меню. "
        "Сохрани порядок строк и колонок. Не добавляй комментарии, верни только текст."
    )
    data_url = f"data:{mime};base64,{b64}"

    resp = client.chat.completions.create(
        model=model,
        messages=[{
            "role": "user",
            "content": [
                {"type": "text", "text": prompt},
                {"type": "image_url", "image_url": {"url": data_url}},
            ],
        }],
        temperature=0.0,
    )
    text = resp.choices[0].message.content or ""
    return text.strip()

# =======================
#   ЧАНКИ/ЭМБЕДДИНГИ
# =======================
def chunk_text(text: str, max_chars: int = 1200, overlap: int = 150) -> List[str]:
    text = " ".join(text.split())
    chunks: List[str] = []
    start = 0
    n = len(text)
    while start < n:
        end = min(n, start + max_chars)
        chunk = text[start:end]
        if chunk.strip():
            chunks.append(chunk)
        if end == n:
            break
        start = max(0, end - overlap)
    return chunks

def embed_texts(client: OpenAI, texts: List[str]) -> List[List[float]]:
    resp = client.embeddings.create(
        model="text-embedding-3-small",
        input=texts,
    )
    return [e.embedding for e in resp.data]

def cosine_sim(a: np.ndarray, b: np.ndarray) -> float:
    na = np.linalg.norm(a); nb = np.linalg.norm(b)
    if na == 0 or nb == 0:
        return 0.0
    return float(np.dot(a, b) / (na * nb))

def retrieve_top_k(question: str, k: int = 4) -> List[Tuple[str, float]]:
    key = _resolve_api_key()
    if not key:
        return []
    client = OpenAI(api_key=key)
    q_emb = client.embeddings.create(
        model="text-embedding-3-small",
        input=[question],
    ).data[0].embedding
    qv = np.array(q_emb, dtype=np.float32)

    con = sqlite3.connect(DB_PATH)
    cur = con.cursor()
    cur.execute("SELECT text, embedding_json FROM chunks")
    rows = cur.fetchall()
    con.close()

    scored: List[Tuple[str, float]] = []
    for t, ej in rows:
        vec = np.array(json.loads(ej), dtype=np.float32)
        scored.append((t, cosine_sim(qv, vec)))
    scored.sort(key=lambda x: x[1], reverse=True)
    return scored[:k]

# =======================
#  НОРМАЛИЗАЦИЯ ПРАЙСА
# =======================
_CURRENCY_MAP = {
    "₸": "KZT", "тг": "KZT", "тенге": "KZT", "kzt": "KZT",
    "₽": "RUB", "руб": "RUB", "руб.": "RUB", "рублей": "RUB", "rub": "RUB",
    "$": "USD", "usd": "USD",
    "€": "EUR", "eur": "EUR",
}

# цена: 2 500, 2.500, 2,500.00, 2500, 2 500 тг, 2500₸, 25.000 руб, 9,90 €
_PRICE_RE = re.compile(
    r"""(?P<num>
            \d{1,3}(?:[ .]\d{3})*(?:[.,]\d{1,2})?   # 1 234,56  | 1.234,56 | 1 234 | 1234.50
            |
            \d+(?:[.,]\d{1,2})?                     # 1234 | 1234,50
        )
        \s*
        (?P<cur>₸|тг|тенге|kzt|₽|руб\.?|рублей|rub|\$|usd|€|eur)? # валюта (опц.)
    """,
    re.IGNORECASE | re.VERBOSE
)

def _normalize_number(num_str: str) -> float:
    s = num_str.replace(" ", "").replace("\u00A0", "")
    # если формат "1.234,56" (евро‑стиль): заменим тысячи '.' и десятые ',' → '.'
    if re.search(r"\d+\.\d{3}(?:\.|\b)", s) and "," in s:
        s = s.replace(".", "").replace(",", ".")
    else:
        # иначе просто заменим запятую на точку для дробной части
        s = s.replace(",", ".")
    try:
        return float(s)
    except Exception:
        # fallback: убираем всё кроме цифр и точки
        s = re.sub(r"[^0-9.]", "", s)
        return float(s) if s else 0.0

def _normalize_currency(cur: Optional[str]) -> Optional[str]:
    if not cur:
        return None
    c = cur.lower().strip().strip(".")
    return _CURRENCY_MAP.get(c, cur.upper())

def extract_catalog_items(text: str) -> List[Dict[str, Any]]:
    """
    Простая эвристика построчного парсинга прайса:
    - Ищем число + опц. валюту
    - Название — это строка без цены или часть до/после цены
    Возвращает список {name, price_value, currency, raw_line, line_no}
    """
    items: List[Dict[str, Any]] = []
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    for i, line in enumerate(lines, start=1):
        m = _PRICE_RE.search(line)
        if not m:
            continue
        num = _normalize_number(m.group("num"))
        cur = _normalize_currency(m.group("cur"))

        # Название товара: берём «всё, кроме найденного прайс‑фрагмента»
        name = (line[:m.start()] + " " + line[m.end():]).strip()
        # Чистим артефакты: пайпы, двойные пробелы
        name = re.sub(r"\s{2,}", " ", name)
        name = name.strip(" |:-—")

        # Если имя пустое — попробуем соседние строки как имя
        if not name and i > 1:
            prev = lines[i - 2]
            if not _PRICE_RE.search(prev) and len(prev) <= 120:
                name = prev

        if not name:
            name = "Позиция"

        items.append({
            "line_no": i,
            "name": name,
            "price_value": num,
            "currency": cur or "KZT",  # по умолчанию KZT, если не указано
            "raw_line": line,
        })
    return items

def db_insert_catalog_items(document_id: int, items: List[Dict[str, Any]]) -> int:
    if not items:
        return 0
    con = sqlite3.connect(DB_PATH)
    cur = con.cursor()
    now = int(time.time())
    for it in items:
        cur.execute("""
            INSERT INTO catalog_items(document_id, line_no, name, price_value, currency, raw_line, created_at)
            VALUES (?, ?, ?, ?, ?, ?, ?)
        """, (
            document_id, it["line_no"], it["name"], it["price_value"], it["currency"], it["raw_line"], now
        ))
    con.commit()
    con.close()
    return len(items)

# =======================
#     ИНДЕКСАЦИЯ
# =======================
def _index_text_blocks(doc_name: str, doc_type: str, local_path: Path, parts: List[str], key: str) -> int:
    client = OpenAI(api_key=key)
    vectors = embed_texts(client, parts)

    con = sqlite3.connect(DB_PATH)
    cur = con.cursor()
    cur.execute(
        "INSERT INTO documents(name, type, path, created_at) VALUES(?,?,?,?)",
        (doc_name, doc_type, str(local_path), int(time.time())),
    )
    doc_id = cur.lastrowid

    for i, (t, vec) in enumerate(zip(parts, vectors)):
        cur.execute(
            "INSERT INTO chunks(document_id, idx, text, embedding_json) VALUES(?,?,?,?)",
            (doc_id, i, t, json.dumps(vec)),
        )

    con.commit()
    con.close()
    return doc_id

# =======================
#  ХЕНДЛЕР ЗАГРУЗКИ
# =======================
async def upload_doc(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    """
    Принимает PDF/DOCX/XLSX/JPG/PNG как документ, парсит, индексирует и вытаскивает прайс‑позиции.
    """
    key = _resolve_api_key()
    if not key:
        await update.message.reply_text(
            "❌ Нет ключа OpenAI. "
            "Добавь свой ключ через /set_openai или укажи платформенный в touch.env (OPENAI_API_KEY=...)."
        )
        return

    msg_doc = update.message.document if update.message else None
    if not msg_doc:
        await update.message.reply_text("Пришли документ (PDF/DOCX/XLSX/JPG/PNG) как файл.")
        return

    fname = (msg_doc.file_name or "").lower()
    supported = any(fname.endswith(ext) for ext in (".pdf", ".docx", ".xlsx", ".jpg", ".jpeg", ".png"))
    if not supported:
        await update.message.reply_text("Сейчас поддерживаем: PDF, DOCX, XLSX, JPG, PNG.")
        return

    await update.message.chat.send_action(ChatAction.UPLOAD_DOCUMENT)

    # Скачиваем
    tg_file = await context.bot.get_file(msg_doc.file_id)
    local_path = UPLOADS_DIR / msg_doc.file_name
    await tg_file.download_to_drive(str(local_path))

    # Парсинг
    try:
        if fname.endswith(".pdf"):
            raw_text = pdf_to_text(str(local_path)); doc_type = "pdf"
        elif fname.endswith(".docx"):
            raw_text = docx_to_text(str(local_path)); doc_type = "docx"
        elif fname.endswith(".xlsx"):
            raw_text = xlsx_to_text(str(local_path)); doc_type = "xlsx"
        elif fname.endswith((".jpg", ".jpeg", ".png")):
            raw_text = image_to_text_openai(str(local_path), api_key=key, model="gpt-4o-mini"); doc_type = "image"
        else:
            await update.message.reply_text("Неподдерживаемый формат."); return

        if not raw_text or not raw_text.strip():
            await update.message.reply_text("Не удалось извлечь текст из файла.")
            return

        # Индексация в эмбеддинги
        parts = chunk_text(raw_text)
        if not parts:
            await update.message.reply_text("Текст извлечён, но пуст после нормализации.")
            return
        doc_id = _index_text_blocks(msg_doc.file_name, doc_type, local_path, parts, key)

        # Извлечение прайс‑позиций
        items = extract_catalog_items(raw_text)
        saved = db_insert_catalog_items(doc_id, items)

        await update.message.reply_text(
            f"✅ Загрузил и проиндексировал файл: {msg_doc.file_name}\n"
            f"Тип: {doc_type.upper()} | Чанков: {len(parts)} | id={doc_id}\n"
            f"🧾 Прайс‑позиций распознано: {saved}"
        )
    except Exception as e:
        await update.message.reply_text(f"Ошибка индексации: {e}")
